"""
Documents Router - Complete Documents Module Implementation
Handles document upload, metadata extraction, search, and file management
"""

from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Query, Form
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, and_, or_, func
from pydantic import BaseModel, Field, validator
from typing import List, Optional, Dict, Any, BinaryIO
from datetime import datetime
import aiofiles
import uuid
import mimetypes
import json
from pathlib import Path
import asyncio
import io

from app.database import get_db
from app.models.document import Document, document_tags
from app.models.tag import Tag
from app.models.user import User
from app.auth.dependencies import get_current_user
from app.config import get_data_dir

router = APIRouter()

# Supported file types and limits
ALLOWED_MIME_TYPES = {
    'application/pdf': '.pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
    'application/msword': '.doc',
    'text/plain': '.txt',
    'image/jpeg': '.jpg',
    'image/png': '.png',
    'image/gif': '.gif',
    'image/webp': '.webp'
}

MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

class DocumentResponse(BaseModel):
    uuid: str
    filename: str
    original_name: str
    mime_type: str
    size_bytes: int
    extracted_text: Optional[str]
    metadata: Dict[str, Any]
    thumbnail_path: Optional[str]
    is_archived: bool
    created_at: datetime
    updated_at: datetime
    tags: List[str]
    
    class Config:
        from_attributes = True

class DocumentSummary(BaseModel):
    uuid: str
    filename: str
    original_name: str
    mime_type: str
    size_bytes: int
    is_archived: bool
    created_at: datetime
    updated_at: datetime
    tags: List[str]
    preview: str  # First 200 chars of extracted text
    
    class Config:
        from_attributes = True

class DocumentUpdate(BaseModel):
    tags: Optional[List[str]] = Field(None, max_items=20)
    is_archived: Optional[bool] = None
    metadata: Optional[Dict[str, Any]] = None

@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    tags: Optional[str] = Form(None),  # JSON string of tag list
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Upload a document with automatic metadata extraction"""
    
    # Validate file
    if not file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No filename provided"
        )
    
    # Check file size
    if hasattr(file, 'size') and file.size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB"
        )
    
    # Validate MIME type
    content_type = file.content_type
    if content_type not in ALLOWED_MIME_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File type not supported. Allowed types: {list(ALLOWED_MIME_TYPES.values())}"
        )
    
    # Generate unique filename
    document_uuid = str(uuid.uuid4())
    file_extension = ALLOWED_MIME_TYPES[content_type]
    filename = f"{document_uuid}{file_extension}"
    
    # Create document directory
    documents_dir = get_data_dir() / "assets" / "documents" / str(datetime.now().year)
    documents_dir.mkdir(parents=True, exist_ok=True)
    
    file_path = documents_dir / filename
    
    # Save file
    content = await file.read()
    file_size = len(content)
    
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB"
        )
    
    async with aiofiles.open(file_path, 'wb') as f:
        await f.write(content)
    
    # Extract text and metadata
    extracted_text = await _extract_text_from_file(file_path, content_type)
    metadata = await _extract_metadata(file_path, content_type, file.filename)
    
    # Generate thumbnail for images
    thumbnail_path = None
    if content_type.startswith('image/'):
        thumbnail_path = await _generate_thumbnail(file_path, document_uuid)
    
    # Parse tags
    tag_list = []
    if tags:
        try:
            tag_list = json.loads(tags)
        except json.JSONDecodeError:
            # Treat as comma-separated string
            tag_list = [tag.strip() for tag in tags.split(',') if tag.strip()]
    
    # Create document record
    document = Document(
        uuid=document_uuid,
        filename=filename,
        original_name=file.filename,
        filepath=str(file_path),
        mime_type=content_type,
        size_bytes=file_size,
        extracted_text=extracted_text,
        metadata_json=json.dumps(metadata),
        thumbnail_path=thumbnail_path,
        user_id=current_user.id
    )
    
    db.add(document)
    await db.flush()
    
    # Handle tags
    await _handle_document_tags(db, document, tag_list, current_user.id)
    
    await db.commit()
    
    return await _get_document_with_relations(db, document_uuid, current_user.id)

@router.get("/{document_uuid}", response_model=DocumentResponse)
async def get_document(
    document_uuid: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Get document metadata and information"""
    
    result = await db.execute(
        select(Document).where(
            and_(Document.uuid == document_uuid, Document.user_id == current_user.id)
        )
    )
    document = result.scalar_one_or_none()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    return await _get_document_with_relations(db, document_uuid, current_user.id)

@router.put("/{document_uuid}", response_model=DocumentResponse)
async def update_document(
    document_uuid: str,
    update_data: DocumentUpdate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Update document metadata and tags"""
    
    result = await db.execute(
        select(Document).where(
            and_(Document.uuid == document_uuid, Document.user_id == current_user.id)
        )
    )
    document = result.scalar_one_or_none()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Update fields
    if update_data.is_archived is not None:
        document.is_archived = update_data.is_archived
    
    if update_data.metadata is not None:
        document.metadata_json = json.dumps(update_data.metadata)
    
    # Handle tags
    if update_data.tags is not None:
        await _handle_document_tags(db, document, update_data.tags, current_user.id)
    
    document.updated_at = datetime.utcnow()
    await db.commit()
    
    return await _get_document_with_relations(db, document_uuid, current_user.id)

@router.delete("/{document_uuid}")
async def delete_document(
    document_uuid: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Delete a document and its file"""
    
    result = await db.execute(
        select(Document).where(
            and_(Document.uuid == document_uuid, Document.user_id == current_user.id)
        )
    )
    document = result.scalar_one_or_none()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Delete physical file
    file_path = Path(document.filepath)
    if file_path.exists():
        file_path.unlink()
    
    # Delete thumbnail if exists
    if document.thumbnail_path:
        thumbnail_path = Path(document.thumbnail_path)
        if thumbnail_path.exists():
            thumbnail_path.unlink()
    
    await db.delete(document)
    await db.commit()
    
    return {"message": "Document deleted successfully"}

@router.get("/{document_uuid}/download")
async def download_document(
    document_uuid: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Download the actual document file"""
    
    result = await db.execute(
        select(Document).where(
            and_(Document.uuid == document_uuid, Document.user_id == current_user.id)
        )
    )
    document = result.scalar_one_or_none()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    file_path = Path(document.filepath)
    if not file_path.exists():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="File not found on disk"
        )
    
    return FileResponse(
        path=file_path,
        filename=document.original_name,
        media_type=document.mime_type
    )

@router.get("/{document_uuid}/preview")
async def preview_document(
    document_uuid: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Get document preview (thumbnail for images, text preview for documents)"""
    
    result = await db.execute(
        select(Document).where(
            and_(Document.uuid == document_uuid, Document.user_id == current_user.id)
        )
    )
    document = result.scalar_one_or_none()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Return thumbnail for images
    if document.thumbnail_path:
        thumbnail_path = Path(document.thumbnail_path)
        if thumbnail_path.exists():
            return FileResponse(
                path=thumbnail_path,
                media_type="image/jpeg"
            )
    
    # Return text preview for text documents
    if document.extracted_text:
        preview_text = document.extracted_text[:1000] + "..." if len(document.extracted_text) > 1000 else document.extracted_text
        return {"preview": preview_text, "type": "text"}
    
    return {"message": "No preview available"}

@router.get("/", response_model=List[DocumentSummary])
async def list_documents(
    mime_type: Optional[str] = Query(None),
    archived: Optional[bool] = Query(False),
    tag: Optional[str] = Query(None),
    search: Optional[str] = Query(None),
    limit: int = Query(50, le=100),
    offset: int = Query(0, ge=0),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """List documents with filtering and search"""
    
    query = select(Document).where(Document.user_id == current_user.id)
    
    # Apply filters
    if mime_type:
        query = query.where(Document.mime_type == mime_type)
    if archived is not None:
        query = query.where(Document.is_archived == archived)
    
    # Tag filter
    if tag:
        query = query.join(document_tags).join(Tag).where(Tag.name == tag)
    
    # Search filter
    if search:
        search_term = f"%{search}%"
        query = query.where(
            or_(
                Document.filename.ilike(search_term),
                Document.original_name.ilike(search_term),
                Document.extracted_text.ilike(search_term)
            )
        )
    
    # Order and pagination
    query = query.order_by(Document.updated_at.desc()).offset(offset).limit(limit)
    
    result = await db.execute(query)
    documents = result.scalars().all()
    
    # Convert to summaries with tags
    summaries = []
    for doc in documents:
        # Get tags
        tag_result = await db.execute(
            select(Tag.name).join(document_tags).where(document_tags.c.document_uuid == doc.uuid)
        )
        doc_tags = [row[0] for row in tag_result.fetchall()]
        
        summaries.append(DocumentSummary(
            uuid=doc.uuid,
            filename=doc.filename,
            original_name=doc.original_name,
            mime_type=doc.mime_type,
            size_bytes=doc.size_bytes,
            is_archived=doc.is_archived,
            created_at=doc.created_at,
            updated_at=doc.updated_at,
            tags=doc_tags,
            preview=doc.extracted_text[:200] + "..." if doc.extracted_text and len(doc.extracted_text) > 200 else (doc.extracted_text or "")
        ))
    
    return summaries

@router.get("/search/fulltext")
async def search_documents(
    query: str = Query(..., min_length=1),
    limit: int = Query(20, le=50),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Full-text search across all document content"""
    
    search_term = f"%{query}%"
    
    # Search in extracted text, filename, and original name
    result = await db.execute(
        select(Document).where(
            and_(
                Document.user_id == current_user.id,
                or_(
                    Document.extracted_text.ilike(search_term),
                    Document.filename.ilike(search_term),
                    Document.original_name.ilike(search_term)
                )
            )
        ).order_by(Document.updated_at.desc()).limit(limit)
    )
    
    documents = result.scalars().all()
    
    # Format results with highlights
    results = []
    for doc in documents:
        # Simple highlight - find matching snippets
        highlight = ""
        if doc.extracted_text:
            text_lower = doc.extracted_text.lower()
            query_lower = query.lower()
            index = text_lower.find(query_lower)
            if index != -1:
                start = max(0, index - 100)
                end = min(len(doc.extracted_text), index + len(query) + 100)
                highlight = doc.extracted_text[start:end]
                if start > 0:
                    highlight = "..." + highlight
                if end < len(doc.extracted_text):
                    highlight = highlight + "..."
        
        results.append({
            "uuid": doc.uuid,
            "original_name": doc.original_name,
            "mime_type": doc.mime_type,
            "highlight": highlight,
            "created_at": doc.created_at
        })
    
    return {"results": results, "total": len(results)}

# Helper functions
async def _handle_document_tags(db: AsyncSession, document: Document, tag_names: List[str], user_id: int):
    """Handle tag assignment for a document"""
    
    # Clear existing tags
    await db.execute(
        document_tags.delete().where(document_tags.c.document_uuid == document.uuid)
    )
    
    for tag_name in tag_names:
        if not tag_name.strip():
            continue
            
        tag_name = tag_name.strip().lower()
        
        # Get or create tag
        result = await db.execute(
            select(Tag).where(
                and_(Tag.name == tag_name, Tag.module_type == "documents", Tag.user_id == user_id)
            )
        )
        tag = result.scalar_one_or_none()
        
        if not tag:
            tag = Tag(
                name=tag_name,
                module_type="documents",
                user_id=user_id
            )
            db.add(tag)
            await db.flush()
        
        # Associate with document
        await db.execute(
            document_tags.insert().values(document_uuid=document.uuid, tag_id=tag.id)
        )

async def _extract_text_from_file(file_path: Path, mime_type: str) -> Optional[str]:
    """Extract text content from uploaded file"""
    
    try:
        if mime_type == 'text/plain':
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                return await f.read()
        
        elif mime_type == 'application/pdf':
            # For PDF extraction, we'd use PyMuPDF
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                return text
            except ImportError:
                return "PDF text extraction not available - PyMuPDF not installed"
        
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            # For DOCX extraction
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except ImportError:
                return "DOCX text extraction not available - python-docx not installed"
        
        return None
        
    except Exception as e:
        print(f"Text extraction failed: {e}")
        return None

async def _extract_metadata(file_path: Path, mime_type: str, original_name: str) -> Dict[str, Any]:
    """Extract metadata from uploaded file"""
    
    metadata = {
        "original_name": original_name,
        "file_size": file_path.stat().st_size,
        "mime_type": mime_type,
        "upload_date": datetime.utcnow().isoformat()
    }
    
    # Add image-specific metadata
    if mime_type.startswith('image/'):
        try:
            from PIL import Image
            with Image.open(file_path) as img:
                metadata.update({
                    "width": img.width,
                    "height": img.height,
                    "format": img.format,
                    "mode": img.mode
                })
        except ImportError:
            pass
        except Exception as e:
            print(f"Image metadata extraction failed: {e}")
    
    return metadata

async def _generate_thumbnail(file_path: Path, document_uuid: str) -> Optional[str]:
    """Generate thumbnail for image files"""
    
    try:
        from PIL import Image
        
        thumbnail_dir = get_data_dir() / "assets" / "images" / "thumbnails"
        thumbnail_dir.mkdir(parents=True, exist_ok=True)
        
        thumbnail_path = thumbnail_dir / f"{document_uuid}_thumb.jpg"
        
        with Image.open(file_path) as img:
            # Convert to RGB if necessary
            if img.mode in ('RGBA', 'LA', 'P'):
                img = img.convert('RGB')
            
            # Create thumbnail
            img.thumbnail((300, 300), Image.Resampling.LANCZOS)
            img.save(thumbnail_path, 'JPEG', quality=85)
        
        return str(thumbnail_path)
        
    except ImportError:
        return None
    except Exception as e:
        print(f"Thumbnail generation failed: {e}")
        return None

async def _get_document_with_relations(db: AsyncSession, document_uuid: str, user_id: int) -> DocumentResponse:
    """Get document with all related data (tags)"""
    
    # Get document
    result = await db.execute(
        select(Document).where(and_(Document.uuid == document_uuid, Document.user_id == user_id))
    )
    document = result.scalar_one()
    
    # Get tags
    tag_result = await db.execute(
        select(Tag.name).join(document_tags).where(document_tags.c.document_uuid == document_uuid)
    )
    tags = [row[0] for row in tag_result.fetchall()]
    
    # Parse metadata
    metadata = {}
    if document.metadata_json:
        try:
            metadata = json.loads(document.metadata_json)
        except json.JSONDecodeError:
            pass
    
    return DocumentResponse(
        uuid=document.uuid,
        filename=document.filename,
        original_name=document.original_name,
        mime_type=document.mime_type,
        size_bytes=document.size_bytes,
        extracted_text=document.extracted_text,
        metadata=metadata,
        thumbnail_path=document.thumbnail_path,
        is_archived=document.is_archived,
        created_at=document.created_at,
        updated_at=document.updated_at,
        tags=tags
    ) 